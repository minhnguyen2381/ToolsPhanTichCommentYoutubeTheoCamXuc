import sys
import pandas as pd
from pathlib import Path
from docx import Document
from paths import DATA_DIR, REPORT_DIR

sys.stdout.reconfigure(encoding='utf-8')

def calculate_stats():
    file_path = DATA_DIR / "v5_3000_videos_filtered.csv"
    if not file_path.exists():
        print(f"File not found: {file_path}")
        return

    df = pd.read_csv(file_path)
    
    # Fill missing with 0 and convert to numeric
    df['views'] = pd.to_numeric(df['views'], errors='coerce').fillna(0)
    df['likes'] = pd.to_numeric(df['likes'], errors='coerce').fillna(0)
    df['comments'] = pd.to_numeric(df['comments'], errors='coerce').fillna(0)
    
    # Sort by views descending and take top 50
    df_top50 = df.sort_values(by='views', ascending=False).head(50)
    
    total_comments = int(df_top50['comments'].sum())
    print(f"Tổng số bình luận của 50 video này là {total_comments} bình luận.\n")
    
    # Bảng 1: Thống kê số lượng bình luận
    comment_bins = [-1, 0, 50, 100, 150, 200, 250, 300, float('inf')]
    comment_labels = ['0', '1-50', '51-100', '101-150', '151-200', '201-250', '251-300', 'Trên 300']
    
    df_top50['comment_group'] = pd.cut(df_top50['comments'], bins=comment_bins, labels=comment_labels)
    comment_stats = df_top50['comment_group'].value_counts().reindex(comment_labels).fillna(0).astype(int)
    
    print("Bảng 1. Thống kê số lượng bình luận của 50 video phổ biến hàng đầu trên YouTube")
    print(f"{'Khoảng số lượng bình luận (bài)':<40} | {'Số lượng video':<15} | {'Tỷ lệ phần trăm':<15}")
    print("-" * 75)
    for label in comment_labels:
        count = comment_stats[label]
        pct = (count / 50) * 100
        print(f"{label:<40} | {count:<15} | {pct:.2f}%")
    print(f"{'Tổng cộng':<40} | {50:<15} | 100%\n")
    
    # Bảng 2: Thống kê số lượng lượt thích
    like_bins = [-1, 500, 1000, 1500, 2000, 2500, 3000, 3500, 4000, 4500, 5000, float('inf')]
    like_labels = ['0-500', '500-1000', '1000-1500', '1500-2000', '2000-2500', '2500-3000', '3000-3500', '3500-4000', '4000-4500', '4500-5000', 'Trên 5000']
    
    df_top50['like_group'] = pd.cut(df_top50['likes'], bins=like_bins, labels=like_labels)
    like_stats = df_top50['like_group'].value_counts().reindex(like_labels).fillna(0).astype(int)
    
    print("Bảng 2. Thống kê số lượng lượt thích của 50 video phổ biến hàng đầu trên YouTube")
    print(f"{'Khoảng số lượng lượt thích (lượt)':<40} | {'Số lượng video':<15} | {'Tỷ lệ phần trăm':<15}")
    print("-" * 75)
    for label in like_labels:
        count = like_stats[label]
        pct = (count / 50) * 100
        print(f"{label:<40} | {count:<15} | {pct:.2f}%")
    print(f"{'Tổng cộng':<40} | {50:<15} | 100%\n")

    # Generate HTML
    html_content = f"""
    <html>
    <head><meta charset="utf-8"><title>Báo cáo phản hồi người dùng</title></head>
    <body style="font-family: Arial, sans-serif; padding: 20px;">
    <h2>3.5.3 Phản hồi của người dùng mạng (bổ sung)</h2>
    <p>Tổng số bình luận của 50 video này là <b>{total_comments}</b> bình luận.</p>
    
    <h3>Bảng 1. Thống kê số lượng bình luận của 50 video phổ biến hàng đầu trên YouTube</h3>
    <table border="1" cellspacing="0" cellpadding="8" style="border-collapse: collapse; width: 100%; max-width: 800px;">
        <tr style="background-color: #f2f2f2;">
            <th align="left">Khoảng số lượng bình luận (bài)</th>
            <th align="center">Số lượng video</th>
            <th align="center">Tỷ lệ phần trăm</th>
        </tr>
    """
    for label in comment_labels:
        count = comment_stats[label]
        pct = (count / 50) * 100
        html_content += f"<tr><td>{label}</td><td align='center'>{count}</td><td align='center'>{pct:.2f}%</td></tr>\n"
    html_content += f"""
        <tr style="font-weight: bold; background-color: #e6e6e6;">
            <td>Tổng cộng</td><td align='center'>50</td><td align='center'>100%</td>
        </tr>
    </table>
    
    <h3>Bảng 2. Thống kê số lượng lượt thích của 50 video phổ biến hàng đầu trên YouTube</h3>
    <table border="1" cellspacing="0" cellpadding="8" style="border-collapse: collapse; width: 100%; max-width: 800px;">
        <tr style="background-color: #f2f2f2;">
            <th align="left">Khoảng số lượng lượt thích (lượt)</th>
            <th align="center">Số lượng video</th>
            <th align="center">Tỷ lệ phần trăm</th>
        </tr>
    """
    for label in like_labels:
        count = like_stats[label]
        pct = (count / 50) * 100
        html_content += f"<tr><td>{label}</td><td align='center'>{count}</td><td align='center'>{pct:.2f}%</td></tr>\n"
    html_content += f"""
        <tr style="font-weight: bold; background-color: #e6e6e6;">
            <td>Tổng cộng</td><td align='center'>50</td><td align='center'>100%</td>
        </tr>
    </table>
    </body></html>
    """
    
    html_path = REPORT_DIR / "vi" / "v8_sentiment_report.html"
    html_path.parent.mkdir(parents=True, exist_ok=True)
    html_path.write_text(html_content, encoding='utf-8')
    print(f"[OK] Đã lưu báo cáo HTML tại {html_path}")
    
    # Generate DOCX
    doc = Document()
    doc.add_heading('3.5.3 Phản hồi của người dùng mạng (bổ sung)', level=2)
    doc.add_paragraph(f"Tổng số bình luận của 50 video này là {total_comments} bình luận.")
    
    doc.add_heading('Bảng 1. Thống kê số lượng bình luận của 50 video phổ biến hàng đầu trên YouTube', level=3)
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Khoảng số lượng bình luận (bài)'
    hdr_cells[1].text = 'Số lượng video'
    hdr_cells[2].text = 'Tỷ lệ phần trăm'
    
    for label in comment_labels:
        row_cells = table1.add_row().cells
        count = comment_stats[label]
        pct = (count / 50) * 100
        row_cells[0].text = label
        row_cells[1].text = str(count)
        row_cells[2].text = f"{pct:.2f}%"
    
    row_cells = table1.add_row().cells
    row_cells[0].text = 'Tổng cộng'
    row_cells[1].text = '50'
    row_cells[2].text = '100%'
    
    doc.add_heading('Bảng 2. Thống kê số lượng lượt thích của 50 video phổ biến hàng đầu trên YouTube', level=3)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Khoảng số lượng lượt thích (lượt)'
    hdr_cells2[1].text = 'Số lượng video'
    hdr_cells2[2].text = 'Tỷ lệ phần trăm'
    
    for label in like_labels:
        row_cells = table2.add_row().cells
        count = like_stats[label]
        pct = (count / 50) * 100
        row_cells[0].text = label
        row_cells[1].text = str(count)
        row_cells[2].text = f"{pct:.2f}%"
        
    row_cells = table2.add_row().cells
    row_cells[0].text = 'Tổng cộng'
    row_cells[1].text = '50'
    row_cells[2].text = '100%'
    
    docx_path = REPORT_DIR / "vi" / "v8_sentiment_report.docx"
    doc.save(str(docx_path))
    print(f"[OK] Đã lưu báo cáo DOCX tại {docx_path}")

if __name__ == "__main__":
    calculate_stats()
