"""BƯỚC 13: Thống kê và Phân tích Người dùng mạng (Audience Stats).
Lấy dữ liệu từ playlist v9, phân loại view, comment, like vào các khoảng
và sinh báo cáo DOCX, HTML với 3 bảng độc lập.
"""

import io
import sys
import os

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from paths import DATA_DIR, REPORT_DIR, ensure_data_dir, ensure_report_dir
from youtube_client import list_playlist_videos, get_video_stats

if sys.stdout.encoding != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

PLAYLIST_ID = "PLlvlc45o3QQcwtas1taX7VlLyZ6Advf_3"
PLAYLIST_CSV = DATA_DIR / "v9_playlist_videos.csv"

# Bins & Labels for View
VIEW_BINS = [-1, 10000, 20000, 30000, 40000, 50000, 60000, 70000, 80000, 90000, 100000, float('inf')]
VIEW_LABELS = [
    "0-10.000",
    "10.000-20.000",
    "20.000-30.000",
    "30.000-40.000",
    "40.000-50.000",
    "50.000-60.000",
    "60.000-70.000",
    "70.000-80.000",
    "80.000-90.000",
    "90.000-100.000",
    "Trên 100.000"
]

# Bins & Labels for Comment
COMMENT_BINS = [-1, 0, 50, 100, 150, 200, 250, 300, float('inf')]
COMMENT_LABELS = [
    "0",
    "1-50",
    "51-100",
    "101-150",
    "151-200",
    "201-250",
    "251-300",
    "300以上"
]

# Bins & Labels for Like
LIKE_BINS = [-1, 500, 1000, 1500, 2000, 2500, 3000, 3500, 4000, 4500, 5000, float('inf')]
LIKE_LABELS = [
    "0-500",
    "500-1000",
    "1000-1500",
    "1500-2000",
    "2000-2500",
    "2500-3000",
    "3000-3500",
    "3500-4000",
    "4000-4500",
    "4500-5000",
    "5000以上"
]

def fetch_playlist_data():
    if PLAYLIST_CSV.exists():
        print(f"[*] Đã thấy {PLAYLIST_CSV.name}, load cache...")
        return pd.read_csv(PLAYLIST_CSV)
    
    print(f"[*] Đang lấy danh sách video từ playlist {PLAYLIST_ID}...")
    video_ids = list_playlist_videos(PLAYLIST_ID)
    print(f"[*] Tìm thấy {len(video_ids)} video trong playlist.")
    
    print("[*] Đang fetch thông tin chi tiết từng video...")
    rows = get_video_stats(video_ids, sleep=0.2)
    df = pd.DataFrame(rows)
    
    # Rename columns to match v5 schema
    if not df.empty:
        df = df.rename(columns={
            "views": "view_count",
            "likes": "like_count",
            "comments": "comment_count"
        })
    
    df.to_csv(PLAYLIST_CSV, index=False, encoding="utf-8-sig")
    print(f"[OK] Đã lưu dữ liệu playlist ra {PLAYLIST_CSV.name}")
    return df

def normalize_column(df, possible_names, default_name):
    for col in possible_names:
        if col in df.columns:
            df[default_name] = pd.to_numeric(df[col], errors='coerce').fillna(0)
            return df
    df[default_name] = 0
    return df

def normalize_metrics(df):
    df = normalize_column(df, ['view_count', 'viewCount', 'views', 'ViewCount'], 'view_count')
    df = normalize_column(df, ['comment_count', 'commentCount', 'comments', 'CommentCount'], 'comment_count')
    df = normalize_column(df, ['like_count', 'likeCount', 'likes', 'LikeCount'], 'like_count')
    return df

def generate_reports(stats_view, stats_comment, stats_like, top_50_count, text_paragraph1, text_paragraph2):
    ensure_report_dir()
    
    # Generate HTML
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset='utf-8'>
        <title>Thống kê người dùng mạng</title>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
            h2 {{ color: #333; }}
            table {{ border-collapse: collapse; width: 100%; max-width: 600px; margin-top: 20px; margin-bottom: 20px; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; font-weight: bold; }}
            .text-content {{ margin-top: 20px; text-indent: 30px; }}
            .bold {{ font-weight: bold; }}
        </style>
    </head>
    <body>
        <h2>3.4.5. Người dùng mạng ( phần 传播受众)</h2>
        <p class="text-content">{text_paragraph1}</p>
        
        <p class="bold">Bảng .... Top 50 video hot Youtube về Quan Vũ</p>
        <table>
            <tr>
                <th>Khoảng lượt xem video phổ biến (lượt)</th>
                <th>Số lượng video</th>
                <th>Tỷ lệ phần trăm</th>
            </tr>
    """
    
    # View Table
    for _, row in stats_view.iterrows():
        lbl = row['Khoảng lượt xem video phổ biến (lượt)'].replace('\n', '<br><br>')
        html_content += f"""
            <tr>
                <td>{lbl}</td>
                <td style="text-align:center;">{row['Số lượng video']}</td>
                <td style="text-align:center;">{row['Tỷ lệ phần trăm']}</td>
            </tr>
        """
    html_content += f"""
            <tr>
                <td class="bold">Tổng cộng</td>
                <td class="bold" style="text-align:center;">{top_50_count}</td>
                <td class="bold" style="text-align:center;">100%</td>
            </tr>
        </table>
        
        <p class="text-content">{text_paragraph2}</p>
        
        <p class="bold">Bảng lượt comment</p>
        <table>
            <tr>
                <th>Comment（条）</th>
                <th>Số lượng（条）</th>
                <th>Phần trăm</th>
            </tr>
    """
    
    # Comment Table
    for _, row in stats_comment.iterrows():
        html_content += f"""
            <tr>
                <td>{row['Comment（条）']}</td>
                <td style="text-align:center;">{row['Số lượng（条）']}</td>
                <td style="text-align:center;">{row['Phần trăm']}</td>
            </tr>
        """
    html_content += f"""
            <tr>
                <td class="bold">Tổng cộng</td>
                <td class="bold" style="text-align:center;">{top_50_count}</td>
                <td class="bold" style="text-align:center;">100%</td>
            </tr>
        </table>
        
        <p class="bold">Bảng lượt like</p>
        <table>
            <tr>
                <th>Like（个）</th>
                <th>Số lượng（个）</th>
                <th>Phần trăm</th>
            </tr>
    """
    
    # Like Table
    for _, row in stats_like.iterrows():
        html_content += f"""
            <tr>
                <td>{row['Like（个）']}</td>
                <td style="text-align:center;">{row['Số lượng（个）']}</td>
                <td style="text-align:center;">{row['Phần trăm']}</td>
            </tr>
        """
    html_content += f"""
            <tr>
                <td class="bold">Tổng cộng</td>
                <td class="bold" style="text-align:center;">{top_50_count}</td>
                <td class="bold" style="text-align:center;">100%</td>
            </tr>
        </table>
    </body>
    </html>
    """
    
    html_path = REPORT_DIR / "vi" / "report_audience.html"
    html_path.parent.mkdir(parents=True, exist_ok=True)
    html_path.write_text(html_content, encoding='utf-8')
    print(f"[OK] Đã lưu báo cáo HTML tại {html_path}")
    
    # Generate DOCX
    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Heading
    doc.add_heading('3.4.5. Người dùng mạng ( phần 传播受众)', level=2)
    
    # Para 1
    doc.add_paragraph(text_paragraph1)
    
    # Table View
    p_title1 = doc.add_paragraph("Bảng .... Top 50 video hot Youtube về Quan Vũ")
    p_title1.runs[0].bold = True
    
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = 'Table Grid'
    h1 = t1.rows[0].cells
    h1[0].text = 'Khoảng lượt xem video phổ biến (lượt)'
    h1[0].paragraphs[0].runs[0].bold = True
    h1[1].text = 'Số lượng video'
    h1[1].paragraphs[0].runs[0].bold = True
    h1[2].text = 'Tỷ lệ phần trăm'
    h1[2].paragraphs[0].runs[0].bold = True
    
    for _, row in stats_view.iterrows():
        r = t1.add_row().cells
        r[0].text = str(row['Khoảng lượt xem video phổ biến (lượt)'])
        r[1].text = str(row['Số lượng video'])
        r[2].text = str(row['Tỷ lệ phần trăm'])
        
    f1 = t1.add_row().cells
    f1[0].text = 'Tổng cộng'
    f1[0].paragraphs[0].runs[0].bold = True
    f1[1].text = str(top_50_count)
    f1[1].paragraphs[0].runs[0].bold = True
    f1[2].text = '100%'
    f1[2].paragraphs[0].runs[0].bold = True
        
    doc.add_paragraph()
    doc.add_paragraph(text_paragraph2)
    
    # Table Comment
    p_title2 = doc.add_paragraph("Bảng lượt comment")
    p_title2.runs[0].bold = True
    
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    h2 = t2.rows[0].cells
    h2[0].text = 'Comment（条）'
    h2[0].paragraphs[0].runs[0].bold = True
    h2[1].text = 'Số lượng（条）'
    h2[1].paragraphs[0].runs[0].bold = True
    h2[2].text = 'Phần trăm'
    h2[2].paragraphs[0].runs[0].bold = True
    
    for _, row in stats_comment.iterrows():
        r = t2.add_row().cells
        r[0].text = str(row['Comment（条）'])
        r[1].text = str(row['Số lượng（条）'])
        r[2].text = str(row['Phần trăm'])

    f2 = t2.add_row().cells
    f2[0].text = 'Tổng cộng'
    f2[0].paragraphs[0].runs[0].bold = True
    f2[1].text = str(top_50_count)
    f2[1].paragraphs[0].runs[0].bold = True
    f2[2].text = '100%'
    f2[2].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    
    # Table Like
    p_title3 = doc.add_paragraph("Bảng lượt like")
    p_title3.runs[0].bold = True
    
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = 'Table Grid'
    h3 = t3.rows[0].cells
    h3[0].text = 'Like（个）'
    h3[0].paragraphs[0].runs[0].bold = True
    h3[1].text = 'Số lượng（个）'
    h3[1].paragraphs[0].runs[0].bold = True
    h3[2].text = 'Phần trăm'
    h3[2].paragraphs[0].runs[0].bold = True
    
    for _, row in stats_like.iterrows():
        r = t3.add_row().cells
        r[0].text = str(row['Like（个）'])
        r[1].text = str(row['Số lượng（个）'])
        r[2].text = str(row['Phần trăm'])
        
    f3 = t3.add_row().cells
    f3[0].text = 'Tổng cộng'
    f3[0].paragraphs[0].runs[0].bold = True
    f3[1].text = str(top_50_count)
    f3[1].paragraphs[0].runs[0].bold = True
    f3[2].text = '100%'
    f3[2].paragraphs[0].runs[0].bold = True
        
    docx_path = REPORT_DIR / "vi" / "report_audience.docx"
    doc.save(str(docx_path))
    print(f"[OK] Đã lưu báo cáo DOCX tại {docx_path}")

def calc_stats(df, col_name, bins, labels, col1_name, col2_name, col3_name):
    # Lấy top 50
    top_50 = df.nlargest(50, col_name).copy()
    if len(top_50) == 0:
        return pd.DataFrame(), 0
        
    # Phân nhóm
    top_50['bin'] = pd.cut(top_50[col_name], bins=bins, labels=labels, right=True)
    
    stats = top_50['bin'].value_counts().reindex(labels, fill_value=0).reset_index()
    stats.columns = [col1_name, col2_name]
    
    stats[col3_name] = (stats[col2_name] / len(top_50) * 100).apply(lambda x: f"{x:.1f}%" if x > 0 else "0%")
        
    return stats, len(top_50)

def main():
    ensure_data_dir()
    df = fetch_playlist_data()
    
    if len(df) == 0:
        print("[!] Playlist rỗng.")
        return

    print(f"[*] Đọc {len(df)} video từ playlist.")
    
    df = normalize_metrics(df)
    
    # View stats
    stats_view, count_view = calc_stats(
        df, 'view_count', VIEW_BINS, VIEW_LABELS, 
        'Khoảng lượt xem video phổ biến (lượt)', 'Số lượng video', 'Tỷ lệ phần trăm'
    )
    
    # Comment stats
    stats_comment, count_comment = calc_stats(
        df, 'comment_count', COMMENT_BINS, COMMENT_LABELS,
        'Comment（条）', 'Số lượng（条）', 'Phần trăm'
    )
    
    # Like stats
    stats_like, count_like = calc_stats(
        df, 'like_count', LIKE_BINS, LIKE_LABELS,
        'Like（个）', 'Số lượng（个）', 'Phần trăm'
    )
    
    if len(stats_view) == 0:
        print("[!] Không có video nào để thống kê.")
        return
        
    # Tính các biến nội suy cho đoạn văn bản
    most_views_bin = stats_view.loc[stats_view['Số lượng video'].idxmax()]
    most_views_range = most_views_bin['Khoảng lượt xem video phổ biến (lượt)'].replace("\n", " hoặc ")
    most_views_pct = most_views_bin['Tỷ lệ phần trăm']
    
    text_p1 = (
        "Trong thời đại Internet hiện nay, các mạng xã hội mới như YouTube, Facebook, Tiktok "
        "đã trở thành kênh chủ yếu để giới trẻ đương đại truyền tải và tiếp nhận thông tin. "
        "Việc thu thập và tổng hợp số lượt xem và số lượng bình luận của các video liên quan "
        "đến Quan Vũ trên các nền tảng mạng có thể giúp chúng ta hiểu được phần nào tổng quy mô "
        "người xem của các video này. Vì vậy, nghiên cứu đã lựa chọn các video liên quan đến "
        "Quan Vũ trên YouTube, đồng thời lọc ra 50 video có mức độ liên quan và lượt xem "
        "xếp hạng cao nhất để tiến hành thống kê. Sau khi tổng hợp dữ liệu về lượt xem "
        "và số lượng bình luận, kết quả cụ thể được thể hiện trong Bảng...."
    )
    
    text_p2 = (
        f"Theo thống kê trong Bảng … các video liên quan đến Quan Vũ có lượng lượt xem tương đối "
        f"lớn trên nền tảng mạng. Lượt xem của các video chủ yếu phân bố trong khoảng {most_views_range}. "
        f"Trong đó, các video có lượt xem từ {most_views_range} chiếm {most_views_pct} tổng số. "
        "Có thể thấy, mặc dù các video liên quan đến Quan Vũ trên các nền tảng mạng xã hội nước ngoài "
        "có số lượng lượt xem tương đối ấn tượng, thu hút được một lượng khán giả nhất định."
    )
    
    print(f"[*] Thống kê Top {count_view} theo View:")
    print(stats_view.to_string(index=False))
    
    # Sinh báo cáo
    generate_reports(stats_view, stats_comment, stats_like, count_view, text_p1, text_p2)
    
    out_csv_view = DATA_DIR / "v9_audience_stats_view.csv"
    stats_view.to_csv(out_csv_view, index=False, encoding='utf-8-sig')
    
    out_csv_comment = DATA_DIR / "v9_audience_stats_comment.csv"
    stats_comment.to_csv(out_csv_comment, index=False, encoding='utf-8-sig')
    
    out_csv_like = DATA_DIR / "v9_audience_stats_like.csv"
    stats_like.to_csv(out_csv_like, index=False, encoding='utf-8-sig')
    
    print("[OK] Đã lưu thống kê raw tại:")
    print(f"  - {out_csv_view.name}")
    print(f"  - {out_csv_comment.name}")
    print(f"  - {out_csv_like.name}")

if __name__ == "__main__":
    main()
